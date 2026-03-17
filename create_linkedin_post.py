"""Generate a LinkedIn post .docx file describing the March Machine Learning Mania 2026 project."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# -- Style defaults --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x1B, 0x1B)
    return h

def add_body(text):
    return doc.add_paragraph(text)

def add_bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

# ============================================================
# DOCUMENT CONTENT
# ============================================================

add_heading('March Machine Learning Mania 2026: Building an NCAA Tournament Prediction Pipeline', level=1)

add_body(
    "March Madness is one of the most exciting — and unpredictable — events in sports. "
    "This year I competed in Kaggle's March Machine Learning Mania 2026 competition, "
    "where the goal is to predict the probability of every possible team matchup in both "
    "the Men's and Women's NCAA basketball tournaments. Your submission is scored by "
    "Brier score — the mean squared error between your predicted probabilities and the "
    "actual binary outcomes. Lower is better, and overconfident wrong predictions are "
    "punished heavily."
)

add_body(
    "Here's a walkthrough of how I built the full pipeline from raw data to final submission."
)

# -- Step 1 --
add_heading('Step 1: Data Collection & Joining', level=2)

add_body(
    "The competition provides decades of NCAA basketball data: game scores, box scores, "
    "tournament seeds, conference affiliations, coaching records, and — for men's only — "
    "weekly rankings from ~192 rating systems (Massey Ordinals). "
    "The men's data spans 41 seasons (1985–2026) with ~199K regular season games and 2,585 "
    "tournament games. The women's data spans 28 seasons (1998–2026) with ~143K regular season "
    "games and 1,717 tournament games."
)

add_body(
    "I merged and aggregated all of these sources into clean per-team, per-season datasets — "
    "computing season averages, shooting percentages, efficiency ratings, Elo ratings, "
    "momentum features (weighted win percentage, win/loss streaks), and strength of schedule. "
    "For men's, I also processed the Massey Ordinals, filtering to pre-tournament rankings "
    "only to prevent data leakage."
)

add_body(
    "Since women's basketball has no Massey Ordinals, I built a synthetic quality ranking "
    "using Ridge regression trained on tournament seeds as a proxy target — essentially "
    "predicting \"how good is this team\" from their season stats."
)

# -- Step 2 --
add_heading('Step 2: Exploratory Data Analysis', level=2)

add_body(
    "EDA revealed several key insights that shaped the modeling approach:"
)

add_bullet(
    "Seeds are the single most predictive feature. 1-seeds win 98.8% of first-round games "
    "against 16-seeds, while 8-vs-9 matchups are essentially coin flips (48.1%)."
)
add_bullet(
    "The women's tournament is significantly more predictable than the men's — "
    "the overall upset rate is 21.1% vs 27.3% for men's."
)
add_bullet(
    "Among Massey Ordinal systems, composites of top systems (KenPom, Sagarin, etc.) "
    "are highly predictive, reaching ~73-78% accuracy on tournament outcomes."
)
add_bullet(
    "Net efficiency (offensive efficiency minus defensive efficiency), win percentage, "
    "and average point differential are the strongest separators between tournament "
    "and non-tournament teams."
)
add_bullet(
    "High multicollinearity between features (WinPct, PointDiff, NetEff are r > 0.94) "
    "meant regularization and feature selection would be important."
)

# -- Step 3 --
add_heading('Step 3: Feature Engineering', level=2)

add_body(
    "For each possible matchup (Team A vs Team B), I computed difference features: "
    "Team A's stat minus Team B's stat. This is the dominant pattern in winning Kaggle solutions "
    "for this competition. The result was 74 features for men's and 63 for women's, spanning:"
)

add_bullet("Tournament seed differences")
add_bullet("Massey/synthetic ranking differences")
add_bullet("Elo rating differences")
add_bullet("Efficiency metrics (offensive, defensive, net)")
add_bullet("Shooting percentages (FG%, 3PT%, FT%, eFG%, TS%)")
add_bullet("Box score margin features (rebounds, assists, turnovers, steals, blocks)")
add_bullet("Momentum features (weighted win %, win/loss streaks)")
add_bullet("Strength of schedule and quality wins against ranked opponents")
add_bullet("Home/road win percentages")
add_bullet("Recent form (last 14 games before tournament)")

add_body(
    "A critical augmentation step was \"flip and double\" — each training matchup appears "
    "twice (original and mirror with negated features and flipped label). This prevents "
    "the model from learning artifacts based on team ID ordering and produces perfectly "
    "balanced 50/50 labels."
)

# -- Step 4 --
add_heading('Step 4: Cross-Validation Strategy', level=2)

add_body(
    "I used Leave-One-Season-Out (LOGO) cross-validation — the standard approach validated "
    "by past winning solutions. Each season serves as one fold: train on all other seasons' "
    "tournament games, predict the held-out season. This gave 40 folds for men's (1985–2025, "
    "excluding the canceled 2020 tournament) and 27 folds for women's (1998–2025)."
)

add_body(
    "This produces honest out-of-fold predictions for every historical tournament game, "
    "which are then used for model comparison and ensemble weight optimization."
)

# -- Step 5 --
add_heading('Step 5: Model Training', level=2)

add_body(
    "I trained three models per gender, each bringing structural diversity to the ensemble:"
)

p = add_body("")
run = p.add_run("Men's ensemble: ")
run.bold = True
p.add_run("XGBoost + PyTorch neural network + Logistic Regression")

p = add_body("")
run = p.add_run("Women's ensemble: ")
run.bold = True
p.add_run("CatBoost + PyTorch neural network + Logistic Regression")

add_body(
    "CatBoost was used for women's instead of XGBoost because it handles missing values "
    "natively — important since ~44% of women's training data lacks detailed box scores (pre-2010)."
)

add_body("Each model went through three phases:")

add_bullet(
    "Phase 1 — Hyperparameter tuning: Optuna Bayesian optimization (30–50 trials) "
    "on the 4 most recent seasons (2022–2025) to find the best configuration."
)
add_bullet(
    "Phase 2 — Full LOGO cross-validation: Using the tuned hyperparameters, generate "
    "out-of-fold predictions across all folds for ensemble weight optimization."
)
add_bullet(
    "Phase 3 — Final model: Train on ALL historical data with the same hyperparameters, "
    "then generate 2026 predictions."
)

add_body(
    "A key modeling decision was training the PyTorch neural network directly on "
    "Brier loss (MSE of predicted probabilities) rather than binary cross-entropy. "
    "This directly optimizes the competition metric and produced better-calibrated "
    "probabilities. The XGBoost model also used a custom Brier loss objective."
)

add_body(
    "After training, isotonic regression calibration was applied to each model's "
    "out-of-fold predictions, improving Brier scores by 0.003–0.004. "
    "Final predictions were clipped to [0.02, 0.98] to avoid extreme overconfidence."
)

# -- Step 6 --
add_heading('Step 6: Ensemble Construction', level=2)

add_body(
    "Ensemble weights were optimized by minimizing Brier score on the calibrated "
    "out-of-fold predictions using scipy's SLSQP constrained optimizer (non-negative "
    "weights summing to 1)."
)

add_body("The optimizer found:")

add_bullet("Men's: PyTorch 57.5% + XGBoost 42.5% + Logistic Regression 0%")
add_bullet("Women's: PyTorch 70.3% + CatBoost 29.7% + Logistic Regression 0%")

add_body(
    "In both cases, logistic regression received zero weight — its predictions were "
    "too correlated with the other models to add ensemble diversity. The PyTorch model "
    "trained on Brier loss dominated both ensembles."
)

# -- Step 7 --
add_heading('Step 7: Results', level=2)

add_body("Final ensemble Brier scores:")

add_bullet("Men's ensemble: 0.1747 (OOF across all historical games)")
add_bullet("Women's ensemble: 0.1299 (OOF across all historical games)")
add_bullet("Men's Stage 1 validation (2022–2025): 0.1801")
add_bullet("Women's Stage 1 validation (2022–2025): 0.1248")

add_body(
    "For context, a naive model predicting 0.5 for every game scores ~0.250, "
    "and a seed-only model scores ~0.210–0.230. Top Kaggle solutions in past years "
    "have scored in the 0.115–0.125 range."
)

add_body(
    "The women's tournament is significantly easier to predict than the men's "
    "(Brier 0.130 vs 0.175), reflecting fewer upsets and more dominant top seeds."
)

# -- Key Takeaways --
add_heading('Key Takeaways', level=2)

add_bullet(
    "Training directly on Brier loss was the single most impactful decision — "
    "the PyTorch model outperformed all gradient boosting models for both genders."
)
add_bullet(
    "Seeds remain king. SeedDiff was the most important feature by a wide margin, "
    "but Massey rankings, Elo ratings, and efficiency metrics added meaningful signal."
)
add_bullet(
    "Proper cross-validation matters. LOGO-CV with season-level folds prevents "
    "temporal leakage and gives honest performance estimates."
)
add_bullet(
    "Simple calibration helps. Isotonic regression on out-of-fold predictions "
    "consistently improved every model."
)
add_bullet(
    "Ensemble diversity is hard to achieve. Both ensembles zeroed out logistic "
    "regression, converging to 2-model blends. When models agree too much, "
    "adding more of them doesn't help."
)

# -- Tech Stack --
add_heading('Tech Stack', level=2)

add_bullet("Python, pandas, NumPy, scikit-learn")
add_bullet("XGBoost, CatBoost, PyTorch")
add_bullet("Optuna (Bayesian hyperparameter optimization)")
add_bullet("AWS SageMaker + S3")
add_bullet("Jupyter notebooks for the full pipeline")

# -- Closing --
add_body("")
add_body(
    "The full pipeline — from raw CSVs to 132,133 matchup predictions — runs end to end "
    "in Jupyter notebooks on SageMaker. Every design decision was informed by researching "
    "past winning solutions, and the code is structured so the entire pipeline can be "
    "re-run when new season data becomes available."
)

add_body(
    "#MarchMadness #MachineLearning #Kaggle #DataScience #NCAA #SportsAnalytics "
    "#PyTorch #XGBoost #DeepLearning"
)

# ============================================================
# SAVE
# ============================================================

output_path = "/Users/aaronengland/Documents/March_Machine_Learning_Mania_2026/march_machine_learning_mania_2026/march_madness_2026_linkedin_post.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
